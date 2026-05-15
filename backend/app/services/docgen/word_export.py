"""Word文档导出 — 严格符合中国法院文书标准格式

排版规范：
- 纸张：A4
- 页边距：上3.7cm 下3.5cm 左2.8cm 右2.6cm
- 标题：方正小标宋简体 22pt（小二号）
- 正文：仿宋_GB2312 16pt（三号），回退仿宋
- 一级标题（##）：黑体 16pt（三号）加粗
- 二级标题（###）：楷体 16pt（三号）加粗
- 行距：固定值28.8pt
- 首行缩进：2字符（约0.74cm）
- 页码：底部居中，首页不同
- 段间距：段前0pt，段后0pt
- 支持水印（草稿/机密）
- 支持文档属性（作者、标题、创建日期）
- 主要章节之间自动分页
"""

import asyncio
import re
import logging
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# 法院标准排版参数
COURT_FONT_SETTINGS = {
    "title_font": "方正小标宋简体",
    "title_font_alt": "FZXiaoBiaoSong-B05S",
    "heading1_font": "黑体",
    "heading1_font_fallback": "SimHei",
    "heading2_font": "楷体",
    "heading2_font_fallback": "KaiTi",
    "body_font": "仿宋_GB2312",
    "body_font_fallback": "仿宋",
    "body_font_fallback2": "FangSong",
    "body_font_en": "Times New Roman",
    "title_size": 22,       # 小二号 = 22pt
    "body_size": 16,        # 三号 = 16pt
    "heading_size": 16,     # 三号 = 16pt
    "line_spacing": 28.8,   # 固定值28.8磅
    "first_line_indent": 0.74,  # 2字符 ≈ 0.74cm（三号字）
    # 页边距：上、右、下、左（cm）
    "margin_top": 3.7,
    "margin_right": 2.6,
    "margin_bottom": 3.5,
    "margin_left": 2.8,
    # 段间距
    "paragraph_space_before": 0,
    "paragraph_space_after": 0,
}

# Safe filename characters for Unicode
_SAFE_FILENAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')

# XML-unsafe characters that can corrupt Word documents
_XML_UNSAFE_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')

# Maximum paragraph length before forcing a page break protection split
_MAX_PARAGRAPH_CHARS = 8000

# Major section markers that should trigger a page break before them
_MAJOR_SECTION_MARKERS = (
    "## 一、", "## 二、", "## 三、", "## 四、", "## 五、",
    "## 六、", "## 七、", "## 八、", "## 九、", "## 十、",
    "## 当事人信息", "## 诉讼请求", "## 事实与理由",
    "## 法律依据", "## 答辩意见", "## 上诉请求", "## 上诉理由",
    "## 反诉请求", "## 代理意见", "## 辩护意见",
)


def _sanitize_xml_text(text: str) -> str:
    """Remove characters that are invalid in XML and would corrupt Word documents."""
    if not text:
        return ""
    # Remove control characters except \t (\x09), \n (\x0a), \r (\x0d)
    cleaned = _XML_UNSAFE_RE.sub("", text)
    return cleaned


def _setup_page(doc: Document):
    """设置页面格式：法院标准A4页边距 + 首页不同页码。"""
    for section in doc.sections:
        section.page_width = Cm(21.0)    # A4
        section.page_height = Cm(29.7)   # A4
        section.top_margin = Cm(COURT_FONT_SETTINGS["margin_top"])
        section.bottom_margin = Cm(COURT_FONT_SETTINGS["margin_bottom"])
        section.left_margin = Cm(COURT_FONT_SETTINGS["margin_left"])
        section.right_margin = Cm(COURT_FONT_SETTINGS["margin_right"])

        # 设置页码（底部居中，首页不同）
        _add_page_number(section)

        # Enable different first page header/footer
        section.different_first_page_header_footer = True


def _set_document_properties(doc: Document, title: str, author: str = "LawMind AI"):
    """设置Word文档属性：标题、作者、创建日期。"""
    core_props = doc.core_properties
    core_props.title = _sanitize_xml_text(title)
    core_props.author = author
    core_props.created = datetime.now()
    core_props.modified = datetime.now()
    core_props.category = "法律文书"
    core_props.comments = "由LawMind AI法律助手平台生成"


def _add_watermark(doc: Document, text: str = "草稿"):
    """添加水印（对角线半透明文字水印）。

    Supported text values: "草稿", "机密", "样本" etc.
    This adds a watermark to the default header of each section.
    """
    if not text or len(text) > 10:
        return

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Add a shape (WordArt) as watermark
        # We use the VML shape approach for watermark
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create the watermark shape using XML
        watermark_xml = f'''<w:r {qn("w:xmlspace")}="preserve"
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office">
            <w:rPr>
                <w:noProof/>
            </w:rPr>
            <w:pict>
                <v:shapetype id="_x0000_t136" coordsize="21600,21600"
                    o:spt="136" adj="10800"
                    path="m@7,l@8,m@5,21600l@6,21600e">
                    <v:formulas>
                        <v:f eqn="sum #0 0 10800"/>
                        <v:f eqn="prod #0 2 1"/>
                        <v:f eqn="sum 21600 0 @1"/>
                        <v:f eqn="sum 0 0 @2"/>
                        <v:f eqn="sum 21600 0 @3"/>
                        <v:f eqn="if @0 @3 0"/>
                        <v:f eqn="if @0 21600 @1"/>
                        <v:f eqn="if @0 0 @2"/>
                        <v:f eqn="if @0 @4 21600"/>
                        <v:f eqn="mid @5 @6"/>
                        <v:f eqn="mid @8 @5"/>
                        <v:f eqn="mid @7 @8"/>
                        <v:f eqn="mid @6 @7"/>
                        <v:f eqn="sum @6 0 @5"/>
                    </v:formulas>
                    <v:path textpathok="t" o:connecttype="custom"
                        o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800"
                        o:connectangles="270,180,90,0"/>
                    <v:textpath on="t" fitshape="t"/>
                    <v:handles>
                        <v:h position="#0,bottomRight" xrange="6629,14971"/>
                    </v:handles>
                    <o:lock v:ext="edit" text="t" shapetype="t"/>
                </v:shapetype>
                <v:shape id="PowerPlusWaterMarkObject"
                    o:spid="_x0000_s2049"
                    type="#_x0000_t136"
                    style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:120pt;rotation:315;z-index:-251658752;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin"
                    o:allowincell="f"
                    fillcolor="#C0C0C0"
                    stroked="f">
                    <v:fill opacity=".25"/>
                    <v:textpath style="font-family:&quot;SimHei&quot;;font-size:1pt"
                        string="{_sanitize_xml_text(text)}"/>
                    <w10:wrap anchorx="margin" anchory="margin"
                        xmlns:w10="urn:schemas-microsoft-com:office:word"/>
                </v:shape>
            </w:pict>
        </w:r>'''

        try:
            from lxml import etree
            # Parse the watermark XML snippet
            # Use a simpler approach: add a diagonal text watermark
            shape_element = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            noProof = OxmlElement("w:noProof")
            rPr.append(noProof)
            shape_element.append(rPr)

            # Build the VML shape using OxmlElement for safety
            pict = OxmlElement("w:pict")

            # Simple text-based watermark using a shape
            shape = OxmlElement("v:shape")
            shape.set("id", "LawMindWatermark")
            shape.set("type", "#_x0000_t136")
            shape.set("style",
                "position:absolute;margin-left:0;margin-top:0;"
                "width:450pt;height:100pt;rotation:315;"
                "z-index:-251658752;"
                "mso-position-horizontal:center;"
                "mso-position-horizontal-relative:margin;"
                "mso-position-vertical:center;"
                "mso-position-vertical-relative:margin")
            shape.set("fillcolor", "#D0D0D0")
            shape.set("stroked", "f")

            fill = OxmlElement("v:fill")
            fill.set("opacity", ".20")
            shape.append(fill)

            textpath = OxmlElement("v:textpath")
            textpath.set("style", 'font-family:"SimHei";font-size:60pt')
            textpath.set("string", _sanitize_xml_text(text))
            shape.append(textpath)

            pict.append(shape)
            shape_element.append(pict)
            paragraph._element.append(shape_element)
        except Exception as e:
            logger.warning("Failed to add watermark (non-critical): %s", e)


def _add_page_number(section):
    """添加页码（底部居中），格式为"— 1 —"。支持首页不同。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 段落格式
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # 添加"— "前缀
    run_prefix = p.add_run("— ")
    run_prefix.font.size = Pt(10)
    run_prefix.font.name = "Times New Roman"

    # 添加页码域代码
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_field = p.add_run()
    run_field._element.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run_instr = p.add_run()
    run_instr._element.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end = p.add_run()
    run_end._element.append(fld_char_end)

    # 添加" —"后缀
    run_suffix = p.add_run(" —")
    run_suffix.font.size = Pt(10)
    run_suffix.font.name = "Times New Roman"

    # First page footer: leave empty (no page number on first page)
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    if not first_footer.paragraphs:
        first_footer.add_paragraph()


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False, east_asia: str | None = None):
    """统一设置run的字体属性，包含fallback字体链。"""
    run.font.size = Pt(size_pt)
    run.font.name = COURT_FONT_SETTINGS["body_font_en"]
    run.bold = bold

    east_asia_font = east_asia or font_name
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    rFonts.set(qn("w:eastAsia"), east_asia_font)
    rFonts.set(qn("w:ascii"), COURT_FONT_SETTINGS["body_font_en"])
    rFonts.set(qn("w:hAnsi"), COURT_FONT_SETTINGS["body_font_en"])
    # Set comprehensive fallback fonts for cross-platform compatibility
    rFonts.set(qn("w:eastAsiaTheme"), "minorEastAsia")
    # Hint attribute for additional fallback
    if not rFonts.get(qn("w:cs")):
        rFonts.set(qn("w:cs"), COURT_FONT_SETTINGS["body_font_en"])


def _set_paragraph_format(p, first_line_indent: bool = True, alignment=None, line_spacing: float | None = None):
    """统一设置段落格式。"""
    ls = line_spacing or COURT_FONT_SETTINGS["line_spacing"]
    p.paragraph_format.line_spacing = Pt(ls)
    p.paragraph_format.space_after = Pt(COURT_FONT_SETTINGS["paragraph_space_after"])
    p.paragraph_format.space_before = Pt(COURT_FONT_SETTINGS["paragraph_space_before"])
    if alignment is not None:
        p.alignment = alignment
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(COURT_FONT_SETTINGS["first_line_indent"])


def _add_page_break(doc: Document):
    """添加分页符。"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def _should_page_break_before(stripped: str) -> bool:
    """Check if a heading line should trigger a page break before it."""
    for marker in _MAJOR_SECTION_MARKERS:
        if stripped.startswith(marker):
            return True
    return False


def _add_title(doc: Document, text: str):
    """添加文书标题：方正小标宋 22pt，居中。"""
    safe_text = _sanitize_xml_text(text)
    if not safe_text:
        safe_text = "文书"

    p = doc.add_paragraph()
    _set_paragraph_format(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    run = p.add_run(safe_text)
    _set_run_font(run, COURT_FONT_SETTINGS["title_font"], COURT_FONT_SETTINGS["title_size"], bold=True,
                  east_asia=COURT_FONT_SETTINGS["title_font"])

    # 标题后空一行
    empty_p = doc.add_paragraph()
    _set_paragraph_format(empty_p, first_line_indent=False)


def _parse_inline_formatting(p, text: str):
    """解析并添加行内格式（加粗）到段落。所有文本经过XML安全处理。"""
    safe_text = _sanitize_xml_text(text)
    if not safe_text:
        return
    parts = safe_text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        is_bold = (i % 2 == 1)
        _set_run_font(run, COURT_FONT_SETTINGS["body_font"], COURT_FONT_SETTINGS["body_size"], bold=is_bold,
                      east_asia=COURT_FONT_SETTINGS["body_font"])


def _split_long_paragraph(text: str, max_chars: int = _MAX_PARAGRAPH_CHARS) -> list[str]:
    """Split a very long paragraph into chunks at sentence boundaries.

    This prevents Word from struggling with extremely long paragraphs
    and provides implicit page-break protection.
    """
    if len(text) <= max_chars:
        return [text]

    chunks: list[str] = []
    remaining = text
    while remaining:
        if len(remaining) <= max_chars:
            chunks.append(remaining)
            break

        # Try to split at a sentence boundary near the limit
        split_pos = max_chars
        # Look for Chinese sentence endings: 。！？；
        for offset in range(min(200, max_chars)):
            pos = max_chars - offset
            if pos > 0 and pos < len(remaining) and remaining[pos] in "。！？；\n":
                split_pos = pos + 1
                break

        chunks.append(remaining[:split_pos])
        remaining = remaining[split_pos:]

    return chunks


def _add_markdown_paragraphs(doc: Document, content: str, first_line_indent: bool = True):
    """将 Markdown 文本解析为 Word 段落，严格遵循法院标准排版。

    解析规则：
    - # → 标题级别（居中，方正小标宋）
    - ## → 一级标题（黑体加粗，不缩进）
    - ### → 二级标题（楷体加粗，不缩进）
    - **粗体** → 加粗
    - 普通段落 → 仿宋，首行缩进2字符
    - 空行 → 空段落（保持行距）

    Handles edge cases:
    - Documents with no headings (all plain text)
    - Tables with inconsistent column counts per row
    - Very long paragraphs (auto-split at sentence boundaries)
    - Special characters that break XML
    """
    lines = content.split("\n")
    i = 0
    in_list = False

    # Track whether any headings were found to handle heading-less documents
    has_heading = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            if in_list:
                in_list = False
            p = doc.add_paragraph()
            _set_paragraph_format(p, first_line_indent=False)
            i += 1
            continue

        # 标题级别（#）— 居中
        if stripped.startswith("### "):
            has_heading = True
            in_list = False
            p = doc.add_paragraph()
            _set_paragraph_format(p, first_line_indent=False)
            _parse_inline_formatting(p, stripped[4:])
            # 设置为楷体
            for run in p.runs:
                _set_run_font(run, COURT_FONT_SETTINGS["heading2_font"],
                              COURT_FONT_SETTINGS["heading_size"], bold=True,
                              east_asia=COURT_FONT_SETTINGS["heading2_font"])
            i += 1
            continue

        if stripped.startswith("## "):
            has_heading = True
            in_list = False
            # Add page break before major sections
            if _should_page_break_before(stripped):
                _add_page_break(doc)
            p = doc.add_paragraph()
            _set_paragraph_format(p, first_line_indent=False)
            _parse_inline_formatting(p, stripped[3:])
            # 设置为黑体
            for run in p.runs:
                _set_run_font(run, COURT_FONT_SETTINGS["heading1_font"],
                              COURT_FONT_SETTINGS["heading_size"], bold=True,
                              east_asia=COURT_FONT_SETTINGS["heading1_font"])
            i += 1
            continue

        if stripped.startswith("# "):
            has_heading = True
            in_list = False
            p = doc.add_paragraph()
            _set_paragraph_format(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _parse_inline_formatting(p, stripped[2:])
            for run in p.runs:
                _set_run_font(run, COURT_FONT_SETTINGS["title_font"],
                              COURT_FONT_SETTINGS["title_size"], bold=True,
                              east_asia=COURT_FONT_SETTINGS["title_font"])
            i += 1
            continue

        # 有序列表（数字. 或 数字、或 数字)）
        ordered_match = re.match(r'^(\d+)[.、)]\s*(.*)', stripped)
        if ordered_match:
            in_list = False
            num = ordered_match.group(1)
            rest = ordered_match.group(2)
            p = doc.add_paragraph()
            indent = first_line_indent
            _set_paragraph_format(p, first_line_indent=indent)
            # 编号部分
            num_run = p.add_run(f"{num}. ")
            _set_run_font(num_run, COURT_FONT_SETTINGS["body_font"], COURT_FONT_SETTINGS["body_size"])
            # 内容部分 — handle very long list items
            for chunk in _split_long_paragraph(rest):
                _parse_inline_formatting(p, chunk)
            i += 1
            continue

        # 无序列表（- 或 *）
        if re.match(r'^[-*]\s', stripped):
            in_list = True
            content_text = stripped[2:]
            p = doc.add_paragraph()
            _set_paragraph_format(p, first_line_indent=first_line_indent)
            bullet_run = p.add_run("• ")
            _set_run_font(bullet_run, COURT_FONT_SETTINGS["body_font"], COURT_FONT_SETTINGS["body_size"])
            _parse_inline_formatting(p, content_text)
            i += 1
            continue

        # 表格行检测（| 分隔）
        if stripped.startswith("|") and "|" in stripped[1:]:
            in_list = False
            # 收集连续表格行
            table_lines: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                stripped_line = lines[i].strip()
                # 跳过分隔行
                if re.match(r'^\|[\s\-:|]+\|$', stripped_line):
                    i += 1
                    continue
                cells = [c.strip() for c in stripped_line.split("|")[1:-1]]
                table_lines.append(cells)
                i += 1

            if table_lines:
                # Determine column count from the first (header) row if possible,
                # otherwise use the max across all rows
                max_cols = max(len(row) for row in table_lines) if table_lines else 0
                if max_cols == 0:
                    i += 1
                    continue

                # Normalize all rows to have exactly max_cols columns,
                # padding short rows with empty strings
                for row_idx in range(len(table_lines)):
                    row = table_lines[row_idx]
                    if len(row) < max_cols:
                        table_lines[row_idx] = row + [""] * (max_cols - len(row))
                    elif len(row) > max_cols:
                        table_lines[row_idx] = row[:max_cols]

                table = doc.add_table(rows=len(table_lines), cols=max_cols)
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_lines):
                    for col_idx in range(max_cols):
                        cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                        cell_text = _sanitize_xml_text(cell_text)
                        cell = table.cell(row_idx, col_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _set_paragraph_format(p, first_line_indent=False)
                        run = p.add_run(cell_text)
                        _set_run_font(run, COURT_FONT_SETTINGS["body_font"],
                                      COURT_FONT_SETTINGS["body_size"],
                                      bold=(row_idx == 0))
            continue

        # 普通段落 — split very long paragraphs for page break protection
        in_list = False
        for chunk in _split_long_paragraph(stripped):
            p = doc.add_paragraph()
            _set_paragraph_format(p, first_line_indent=first_line_indent)
            _parse_inline_formatting(p, chunk)
        i += 1


async def export_to_docx(doc_record, output_dir: Path, watermark: str | None = None) -> str:
    """将文书导出为符合法院格式要求的Word文档

    Args:
        doc_record: Document record with title, content, etc.
        output_dir: Directory to save the file
        watermark: Optional watermark text (e.g. "草稿", "机密")
    """
    doc = Document()
    _setup_page(doc)
    _set_document_properties(doc, doc_record.title)
    if watermark:
        _add_watermark(doc, watermark)
    _add_title(doc, doc_record.title)

    # Handle empty or missing content gracefully
    content = getattr(doc_record, 'content', '') or ''
    if not content.strip():
        content = "（文档内容为空）"

    _add_markdown_paragraphs(doc, content, first_line_indent=True)

    safe_title = _SAFE_FILENAME_RE.sub("", doc_record.title)
    if not safe_title:
        safe_title = f"document_{doc_record.id}"
    filename = f"{doc_record.id}_{safe_title}.docx"
    filepath = output_dir / filename

    output_dir.mkdir(parents=True, exist_ok=True)
    await asyncio.to_thread(doc.save, str(filepath))
    return str(filepath)


async def export_research_to_docx(report_record, output_dir: Path, watermark: str | None = None) -> str:
    """将研究报告导出为Word文档

    Args:
        report_record: Research report record
        output_dir: Directory to save the file
        watermark: Optional watermark text (e.g. "草稿", "机密")
    """
    doc = Document()
    _setup_page(doc)
    _set_document_properties(doc, f"法律研究报告：{report_record.query}")
    if watermark:
        _add_watermark(doc, watermark)
    _add_title(doc, f"法律研究报告：{report_record.query}")

    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sources = getattr(report_record, 'sources_used', []) or []
    created = getattr(report_record, 'created_at', None)
    created_str = created.strftime('%Y-%m-%d %H:%M') if created else '未知时间'
    run = meta.add_run(f"生成时间：{created_str}　来源：{'、'.join(sources)}")
    run.font.size = Pt(12)
    run.font.name = COURT_FONT_SETTINGS["body_font_fallback"]
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), COURT_FONT_SETTINGS["body_font"])
    run.font.color.rgb = RGBColor(128, 128, 128)
    _set_paragraph_format(meta, first_line_indent=False)

    # Handle empty report gracefully
    report_text = getattr(report_record, 'report', '') or ''
    if not report_text.strip():
        report_text = "（报告内容为空）"

    _add_markdown_paragraphs(doc, report_text, first_line_indent=False)

    safe_query = _SAFE_FILENAME_RE.sub("", report_record.query[:30])
    if not safe_query:
        safe_query = f"report_{report_record.id}"
    filename = f"research_{report_record.id}_{safe_query}.docx"
    filepath = output_dir / filename

    output_dir.mkdir(parents=True, exist_ok=True)
    await asyncio.to_thread(doc.save, str(filepath))
    return str(filepath)
