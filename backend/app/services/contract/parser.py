"""合同文档解析服务 — 将PDF/Word/图片转为结构化条款

支持：
- PDF（含扫描件OCR识别）
- Word文档（.docx/.doc）
- 纯文本（.txt）
- 图片（.png/.jpg/.jpeg/.bmp/.tiff）— 通过LLM视觉能力OCR
- 大型合同自动分块处理
"""

import asyncio
import base64
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".png", ".jpg", ".jpeg", ".bmp", ".tiff"}

# 41种合同条款类型（参考CUAD分类）
CLAUSE_TYPES = [
    "当事人信息", "标的", "数量", "质量", "价款或报酬",
    "履行期限地点方式", "违约责任", "争议解决", "合同生效", "合同解除",
    "保密条款", "知识产权", "担保条款", "不可抗力", "通知条款",
    "转让条款", "适用法律", "合同期限", "验收条款", "售后服务",
    "付款方式", "交付条款", "所有权转移", "风险转移", "保险条款",
    "税务条款", "审计条款", "竞业限制", "排他条款", "授权许可",
    "归属条款", "补偿条款", "终止条款", "续约条款", "附件条款",
    "定义条款", "陈述保证", "合规条款", "数据保护", "反贿赂",
    "其他条款",
]

# Large document threshold: pages above this trigger chunked processing
_LARGE_DOC_PAGES = 100

# Maximum characters per chunk for large document processing
_CHUNK_SIZE = 15000


def validate_contract_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


async def parse_contract_document(file_path: Path, llm_client=None, model: str = "glm-5.1") -> dict:
    """解析合同文件，返回 {text, clauses}"""
    if not file_path.exists():
        return {"text": "[文件不存在]", "clauses": []}
    if file_path.stat().st_size == 0:
        return {"text": "[文件为空]", "clauses": []}

    text = await _extract_text(file_path, llm_client, model)

    # Handle PDF with no extractable text — attempt OCR via LLM vision
    if text and text.strip() == "[PDF为扫描件，需要OCR识别]":
        logger.info("PDF has no extractable text, attempting OCR via LLM vision")
        ocr_text = await _extract_pdf_ocr(file_path, llm_client, model)
        if ocr_text and len(ocr_text.strip()) > 10:
            text = ocr_text

    if not text or len(text.strip()) < 10:
        # Final attempt: return whatever text we have with empty clauses
        return {"text": text or "[无法提取文本内容]", "clauses": []}

    clauses = await _identify_clauses(text, llm_client, model)
    return {"text": text, "clauses": clauses}


async def _extract_text(file_path: Path, llm_client=None, model: str = "glm-5.1") -> str:
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".pdf":
            return await _extract_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return await _extract_docx(file_path)
        elif suffix == ".txt":
            return file_path.read_text(encoding="utf-8", errors="ignore")
        elif suffix in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
            return await _extract_image(file_path, llm_client, model)
        return ""
    except Exception as e:
        logger.warning(f"Contract text extraction failed for {file_path}: {e}")
        return f"[文本提取失败: {e}]"


async def _extract_pdf(file_path: Path) -> str:
    """Extract text from PDF. Returns special marker for scanned PDFs."""
    def _sync():
        import pypdf
        reader = pypdf.PdfReader(str(file_path))
        page_count = len(reader.pages)
        texts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                texts.append(t.strip())
        return "\n\n".join(texts) if texts else "[PDF为扫描件，需要OCR识别]", page_count
    result, page_count = await asyncio.to_thread(_sync)

    # For very large documents (>100 pages), truncate to avoid overwhelming LLM
    if page_count > _LARGE_DOC_PAGES:
        logger.warning(f"Large PDF detected ({page_count} pages), will use chunked processing")
        result = _truncate_to_chunk_size(result, _CHUNK_SIZE)

    return result


async def _extract_pdf_ocr(file_path: Path, llm_client=None, model: str = "glm-5.1") -> str:
    """Attempt to OCR a scanned PDF using LLM vision capabilities.

    Converts PDF pages to images and sends them to the LLM for text extraction.
    """
    if llm_client is None:
        return ""

    try:
        def _render_pages_to_images():
            """Convert PDF pages to PNG images using pypdf + Pillow."""
            import pypdf
            from PIL import Image
            import io

            reader = pypdf.PdfReader(str(file_path))
            images = []
            for page_num, page in enumerate(reader.pages):
                # Try to extract images from the page
                if hasattr(page, 'images'):
                    for img in page.images:
                        try:
                            pil_img = Image.open(io.BytesIO(img.data))
                            if pil_img.mode == 'RGBA':
                                pil_img = pil_img.convert('RGB')
                            buf = io.BytesIO()
                            pil_img.save(buf, format='JPEG', quality=85)
                            images.append(buf.getvalue())
                            # Limit to first 10 pages for OCR
                            if len(images) >= 10:
                                return images
                        except Exception:
                            continue
            return images

        images = await asyncio.to_thread(_render_pages_to_images)
        if not images:
            return ""

        all_text: list[str] = []
        for img_data in images:
            b64 = base64.b64encode(img_data).decode("utf-8")
            try:
                response = await llm_client.messages.create(
                    model=model,
                    max_tokens=4096,
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": b64}},
                            {"type": "text", "text": "请识别并提取这份合同图片中的所有文字内容，保持原始格式和段落结构。"},
                        ],
                    }],
                )
                page_text = response.content[0].text if response.content else ""
                if page_text:
                    all_text.append(page_text)
            except Exception as e:
                logger.warning(f"PDF page OCR failed: {e}")
                continue

        return "\n\n".join(all_text) if all_text else ""

    except Exception as e:
        logger.warning(f"PDF OCR fallback failed: {e}")
        return ""


def _truncate_to_chunk_size(text: str, max_chars: int) -> str:
    """Truncate text to max_chars, trying to break at paragraph boundaries."""
    if len(text) <= max_chars:
        return text
    # Try to find a good break point near the limit
    truncated = text[:max_chars]
    # Look for the last paragraph break
    last_break = truncated.rfind("\n\n")
    if last_break > max_chars * 0.5:
        return truncated[:last_break] + f"\n\n[注：文档过长（{len(text)}字符），已截取前{last_break}字符进行分析]"
    return truncated + f"\n\n[注：文档过长（{len(text)}字符），已截取前{max_chars}字符进行分析]"


async def _extract_docx(file_path: Path) -> str:
    def _sync():
        from docx import Document
        doc = Document(str(file_path))
        parts = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        return "\n".join(parts)
    return await asyncio.to_thread(_sync)


async def _extract_image(file_path: Path, llm_client, model: str) -> str:
    """Extract text from an image using LLM vision (OCR)."""
    if llm_client is None:
        return "[图片合同需要配置LLM进行OCR识别]"

    image_data = file_path.read_bytes()
    b64 = base64.b64encode(image_data).decode("utf-8")
    suffix = file_path.suffix.lower()
    media_type = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "bmp": "image/bmp", "tiff": "image/tiff"}.get(suffix.lstrip("."), "image/png")

    try:
        response = await llm_client.messages.create(
            model=model,
            max_tokens=8192,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                    {"type": "text", "text": "请识别并提取这份合同图片中的所有文字内容，保持原始格式和段落结构。"},
                ],
            }],
        )
        return response.content[0].text if response.content else "[未能识别文字]"
    except Exception as e:
        logger.warning(f"Contract image OCR failed: {e}")
        return f"[图片文字提取失败: {e}]"


async def _identify_clauses(text: str, llm_client, model: str) -> list[dict]:
    """使用AI识别合同条款类型和内容。

    For very large contracts, chunks the text and processes each chunk.
    """
    if llm_client is None:
        return _split_by_headers(text)

    # For large documents, process in chunks
    if len(text) > _CHUNK_SIZE:
        return await _identify_clauses_chunked(text, llm_client, model)

    return await _identify_clauses_single(text, llm_client, model)


async def _identify_clauses_single(text: str, llm_client, model: str) -> list[dict]:
    """Identify clauses from a single text block."""
    prompt = f"""请分析以下合同文本，识别并分类每个条款。

将合同按条款拆分，每个条款标注类型。支持的条款类型包括：
{', '.join(CLAUSE_TYPES)}

请以JSON数组格式返回，每个元素包含：
- type: 条款类型（从上述列表中选择最匹配的）
- text: 条款原文（保持原文不变）
- position: 条款在文档中的顺序编号（从1开始）

合同文本：
{text[:15000]}

请直接返回JSON数组，不要包含其他文字。"""

    try:
        response = await llm_client.messages.create(
            model=model,
            max_tokens=8192,
            messages=[{"role": "user", "content": prompt}],
        )
        import json
        raw = response.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        clauses = json.loads(raw)
        return clauses if isinstance(clauses, list) else []
    except Exception as e:
        logger.warning(f"AI clause identification failed: {e}")
        return _split_by_headers(text)


async def _identify_clauses_chunked(text: str, llm_client, model: str) -> list[dict]:
    """Process large contracts by splitting into chunks and identifying clauses in each."""
    chunks = _split_text_into_chunks(text, _CHUNK_SIZE)
    all_clauses: list[dict] = []
    position_offset = 0

    for chunk_idx, chunk in enumerate(chunks):
        chunk_clauses = await _identify_clauses_single(chunk, llm_client, model)
        for clause in chunk_clauses:
            clause["position"] = clause.get("position", 1) + position_offset
            all_clauses.append(clause)
        position_offset += len(chunk_clauses)

    logger.info(f"Chunked clause identification: {len(chunks)} chunks, {len(all_clauses)} total clauses")
    return all_clauses


def _split_text_into_chunks(text: str, max_chars: int) -> list[str]:
    """Split text into chunks at paragraph boundaries, each under max_chars."""
    if len(text) <= max_chars:
        return [text]

    chunks: list[str] = []
    paragraphs = text.split("\n\n")
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 > max_chars and current_chunk:
            chunks.append(current_chunk.strip())
            current_chunk = para
        else:
            current_chunk += "\n\n" + para if current_chunk else para

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks


def _split_by_headers(text: str) -> list[dict]:
    """简单的基于标题模式的条款拆分（备用方案）"""
    import re
    clauses = []
    pattern = re.compile(r'(?:^|\n)\s*((?:第[一二三四五六七八九十百千]+条|[一二三四五六七八九十]+[、.．]|\d+[、.．]|【[^】]+】)\s*.+?)(?=\n(?:第[一二三四五六七八九十百千]+条|[一二三四五六七八九十]+[、.．]|\d+[、.．]|【[^】]+】)|$)', re.DOTALL)

    matches = pattern.findall(text)
    if matches:
        for i, m in enumerate(matches):
            lines = m.strip().split("\n", 1)
            header = lines[0].strip()
            body = lines[1].strip() if len(lines) > 1 else ""
            clauses.append({
                "type": _guess_clause_type(header + body),
                "text": m.strip(),
                "position": i + 1,
            })
    else:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for i, p in enumerate(paragraphs):
            clauses.append({
                "type": "其他条款",
                "text": p,
                "position": i + 1,
            })
    return clauses


def _guess_clause_type(text: str) -> str:
    """根据关键词猜测条款类型"""
    keywords_map = {
        "违约责任": ["违约", "赔偿", "罚金", "滞纳金", "违约金"],
        "争议解决": ["仲裁", "诉讼", "管辖", "争议解决", "法院"],
        "保密条款": ["保密", "秘密", "机密", "不公开"],
        "不可抗力": ["不可抗力", "不可预见", "不能避免"],
        "付款方式": ["付款", "支付", "结算", "费用", "价款"],
        "合同期限": ["期限", "有效期", "起止", "届满"],
        "合同生效": ["生效", "签署", "签字", "盖章"],
        "知识产权": ["知识产权", "专利", "商标", "著作权", "版权"],
        "担保条款": ["担保", "保证", "抵押", "质押"],
        "适用法律": ["适用法律", "法律适用"],
        "合同解除": ["解除", "终止", "撤销", "退出"],
        "通知条款": ["通知", "送达", "告知"],
        "竞业限制": ["竞业", "禁止从事"],
        "数据保护": ["数据保护", "个人信息", "隐私"],
    }
    for clause_type, keywords in keywords_map.items():
        if any(kw in text for kw in keywords):
            return clause_type
    return "其他条款"
