"""合同审查路由"""

import asyncio
import uuid
import logging
import tempfile
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import JSONResponse
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.security import get_current_user, validate_upload
from app.config import get_settings
from app.models.user import User
from app.models.case import Case
from app.models.contract import Contract
from app.schemas.contract import ContractOut, ReviewReportExport
from app.services.contract.parser import parse_contract_document
from app.services.contract.engine import review_contract

logger = logging.getLogger(__name__)
router = APIRouter()

CONTRACT_ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".png", ".jpg", ".jpeg", ".bmp", ".tiff"}


def _to_out(row: Contract, truncate: bool = False) -> dict:
    out = ContractOut.model_validate(row)
    out.has_file = row.file_path is not None
    if truncate:
        if out.parsed_text and len(out.parsed_text) > 300:
            out.parsed_text = out.parsed_text[:300] + "..."
        if out.review_report and len(out.review_report) > 300:
            out.review_report = out.review_report[:300] + "..."
    return out.model_dump(mode="json")


@router.get("", response_model=list[ContractOut])
async def list_contracts(
    case_id: int | None = None,
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=100),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    try:
        base_where = Contract.owner_id == current_user.id
        # Count query
        count_q = select(func.count(Contract.id)).where(base_where)
        if case_id:
            count_q = count_q.where(Contract.case_id == case_id)
        total = (await db.execute(count_q)).scalar() or 0

        query = select(Contract).where(base_where)
        if case_id:
            query = query.where(Contract.case_id == case_id)
        query = query.order_by(Contract.created_at.desc()).offset(skip).limit(limit)
        result = await db.execute(query)
        items = [_to_out(c, truncate=True) for c in result.scalars().all()]
        return JSONResponse(
            content=items,
            headers={"X-Total-Count": str(total)},
        )
    except Exception as e:
        logger.error(f"List contracts failed: {e}")
        raise HTTPException(500, "查询合同列表失败")


@router.post("/upload", response_model=ContractOut, status_code=201)
async def upload_contract(
    file: UploadFile = File(...),
    title: str = Form(...),
    case_id: int | None = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    try:
        if not title or not title.strip():
            raise HTTPException(400, "合同标题不能为空")

        # Centralized upload validation with extension allowlist
        content = await validate_upload(file, allowed_extensions=CONTRACT_ALLOWED_EXTENSIONS)

        if case_id:
            case = await db.execute(select(Case).where(Case.id == case_id, Case.owner_id == current_user.id))
            if not case.scalar_one_or_none():
                raise HTTPException(404, "案件不存在")

        settings = get_settings()
        upload_dir = settings.upload_path / "contracts"
        upload_dir.mkdir(parents=True, exist_ok=True)

        ext = Path(file.filename or ".bin").suffix
        safe_name = f"{uuid.uuid4().hex[:12]}{ext}"
        dest = upload_dir / safe_name
        await asyncio.to_thread(dest.write_bytes, content)

        row = Contract(
            owner_id=current_user.id,
            case_id=case_id,
            title=title.strip()[:200],
            file_path=f"contracts/{safe_name}",
            file_type=ext.lstrip("."),
            status="pending",
        )
        db.add(row)
        await db.commit()
        await db.refresh(row)

        # Auto-parse
        try:
            row.status = "parsing"
            await db.commit()

            from app.services.llm_client import create_llm_client_from_settings
            client = create_llm_client_from_settings(settings)
            parsed = await parse_contract_document(dest, client, settings.CLAUDE_MODEL)
            row.parsed_text = parsed["text"]
            row.clauses = parsed["clauses"]
            row.status = "pending"
            await db.commit()
            await db.refresh(row)
        except Exception as e:
            logger.warning(f"Auto-parse failed for contract {row.id}: {e}")
            row.status = "pending"
            await db.commit()
            await db.refresh(row)

        return _to_out(row)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload contract failed: {e}")
        await db.rollback()
        raise HTTPException(500, "上传合同失败")


@router.post("/draft", response_model=ContractOut, status_code=201)
async def draft_contract(
    title: str = Form(...),
    description: str = Form(...),
    case_id: int | None = Form(None),
    file: UploadFile | None = File(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """根据用户描述和可选的参考文件起草合同。"""
    settings = get_settings()
    reference_text = ""

    if file:
        content = await validate_upload(file)
        tmp_dir = settings.upload_path / "tmp"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        ext = Path(file.filename or ".bin").suffix
        tmp_path = tmp_dir / f"{uuid.uuid4().hex[:12]}{ext}"

        try:
            await asyncio.to_thread(tmp_path.write_bytes, content)
            from app.services.evidence.ocr import extract_text
            reference_text = await extract_text(tmp_path)
        finally:
            await asyncio.to_thread(tmp_path.unlink, missing_ok=True)

    # 构建提示词
    prompt_parts = [
        "你是一位专业的合同起草律师。请根据以下需求起草一份完整的合同：\n",
        f"合同标题：{title}\n",
        f"需求描述：{description}\n",
    ]
    if reference_text and not reference_text.startswith("["):
        prompt_parts.append(f"\n参考文件内容：\n{reference_text[:8000]}\n")
    prompt_parts.append("\n请起草一份格式规范、条款完整的合同文本。使用中文，包含所有必要条款（标的、数量、价款、履行方式、违约责任、争议解决等）。")

    try:
        from app.services.llm_client import create_llm_client_from_settings
        client = create_llm_client_from_settings(settings)
        response = await client.messages.create(
            model=settings.CLAUDE_MODEL,
            max_tokens=settings.CLAUDE_MAX_TOKENS,
            messages=[{"role": "user", "content": "\n".join(prompt_parts)}],
        )
        draft_text = response.content[0].text if response.content else ""
    except Exception as e:
        raise HTTPException(503, f"合同起草失败: {str(e)[:200]}")

    row = Contract(
        owner_id=current_user.id,
        case_id=case_id,
        title=title.strip()[:200],
        parsed_text=draft_text,
        file_type="draft",
        status="completed",
    )
    db.add(row)
    await db.commit()
    await db.refresh(row)
    return _to_out(row)


@router.get("/{contract_id}", response_model=ContractOut)
async def get_contract(
    contract_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    try:
        result = await db.execute(
            select(Contract).where(Contract.id == contract_id, Contract.owner_id == current_user.id)
        )
        row = result.scalar_one_or_none()
        if not row:
            raise HTTPException(404, "合同不存在")
        return _to_out(row)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get contract failed: {e}")
        raise HTTPException(500, "查询合同失败")


@router.post("/{contract_id}/review", response_model=ContractOut)
async def review_contract_endpoint(
    contract_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Contract).where(Contract.id == contract_id, Contract.owner_id == current_user.id)
    )
    row = result.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "合同不存在")

    # Ensure parsed text exists
    if not row.parsed_text:
        settings = get_settings()
        if row.file_path:
            fp = settings.upload_path / row.file_path
            if fp.exists():
                from app.services.llm_client import create_llm_client_from_settings
                client = create_llm_client_from_settings(settings)
                parsed = await parse_contract_document(fp, client, settings.CLAUDE_MODEL)
                row.parsed_text = parsed["text"]
                row.clauses = parsed["clauses"]
                await db.commit()
                await db.refresh(row)

    if not row.parsed_text:
        raise HTTPException(400, "合同文本为空，无法审查")

    # Get case context — verify ownership
    case_context = ""
    if row.case_id:
        case_result = await db.execute(
            select(Case).where(Case.id == row.case_id, Case.owner_id == current_user.id)
        )
        case = case_result.scalar_one_or_none()
        if case and case.description:
            case_context = case.description

    row.status = "reviewing"
    await db.commit()

    try:
        review_result = await review_contract(row.parsed_text, row.clauses, case_context)
        row.review_report = review_result["report"]
        row.risk_items = review_result["risk_items"]
        row.risk_score = review_result["risk_score"]
        row.status = "completed"
    except Exception as e:
        logger.error(f"Contract review failed for {contract_id}: {e}")
        row.status = "failed"
        row.review_report = "审查失败，请稍后重试。"

    await db.commit()
    await db.refresh(row)
    return _to_out(row)


@router.post("/{contract_id}/export")
async def export_report(
    contract_id: int,
    data: ReviewReportExport,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Contract).where(Contract.id == contract_id, Contract.owner_id == current_user.id)
    )
    row = result.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "合同不存在")
    if not row.review_report:
        raise HTTPException(400, "尚未完成审查，无报告可导出")

    if data.format == "docx":
        return await _export_docx(row)
    else:
        from fastapi.responses import PlainTextResponse
        return PlainTextResponse(row.review_report, media_type="text/markdown")


async def _export_docx(row: Contract) -> "FileResponse":
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import FileResponse

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.right_margin = Cm(2.8)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.6)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"合同审查报告 — {row.title}")
    run.font.size = Pt(22)
    run.bold = True

    lines = (row.review_report or "").split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            p = doc.add_paragraph()
        elif line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            r.font.size = Pt(18)
            r.bold = True
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:])
            r.font.size = Pt(16)
            r.bold = True
        elif line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:])
            r.font.size = Pt(15)
            r.bold = True
        elif line.startswith("> "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.font.size = Pt(12)
            r.italic = True
        elif line.startswith("- "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.font.size = Pt(14)
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(14)

        p.paragraph_format.line_spacing = Pt(28)

    # Write to temp file to avoid concurrent write issues
    safe_title = "".join(c for c in row.title if c.isalnum() or c in "（）()—")[:50]
    filename = f"审查报告_{safe_title}.docx"
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix=f"contract_{row.id}_")
    await asyncio.to_thread(doc.save, tmp.name)
    tmp.close()

    return FileResponse(
        tmp.name,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.delete("/{contract_id}")
async def delete_contract(
    contract_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    try:
        result = await db.execute(
            select(Contract).where(Contract.id == contract_id, Contract.owner_id == current_user.id)
        )
        row = result.scalar_one_or_none()
        if not row:
            raise HTTPException(404, "合同不存在")

        if row.file_path:
            settings = get_settings()
            fp = settings.upload_path / row.file_path
            if fp.exists():
                fp.unlink()

        await db.delete(row)
        await db.commit()
        return {"message": "已删除"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete contract failed: {e}")
        await db.rollback()
        raise HTTPException(500, "删除合同失败")
